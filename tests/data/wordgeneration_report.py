from docx import Document

# Create the Word document
doc = Document()
doc.add_heading("Project III | Business Case: Building a Multimodal AI ChatBot for YouTube Video QA", 0)

# --- Business Case ---
doc.add_heading("Business Case", level=1)
doc.add_paragraph(
    "Building a chatbot that can translate YouTube videos into text and allow for natural language querying "
    "offers multiple benefits:\n"
    "- Accessibility: helps hearing-impaired users and readers.\n"
    "- Content Search: enables indexing and quick search within video content.\n"
    "- Customer Support: leverages video knowledge for instant answers, reducing costs.\n"
    "- Education: enhances learning by retrieving key parts of videos.\n"
    "- SEO: transcripts improve video discoverability.\n"
    "- Multilingual: caters to global audiences with translations."
)

# --- Project Overview ---
doc.add_heading("Project Overview", level=1)
doc.add_paragraph(
    "The goal is to build a Retrieval-Augmented Generation (RAG) chatbot that combines NLP and speech recognition "
    "to answer questions about YouTube videos. It retrieves video content, transcribes audio, stores text in a "
    "vector database, and uses LangChain agents with memory and tools for interaction."
)

# --- Key Objectives ---
doc.add_heading("Key Objectives", level=1)
doc.add_paragraph(
    "- Develop a text-based QA model using pre-trained LLMs.\n"
    "- Integrate Whisper speech recognition for video/audio input.\n"
    "- Provide a text and optional voice conversational interface.\n"
    "- Retrieve, analyze, and store YouTube transcripts into Chroma vector database.\n"
    "- Use LangChain agents and tools for retrieval, summarization, translation, and nutrition estimation.\n"
    "- Deploy as a web/mobile app with Streamlit."
)

# --- Project Timeline ---
doc.add_heading("Project Timeline", level=1)
doc.add_paragraph(
    "Day 1-2: Data collection and text preprocessing.\n"
    "Day 3-4: QA model development and speech recognition.\n"
    "Day 5-6: Conversational interface and video retrieval.\n"
    "Day 7: Testing, evaluation, documentation, and presentation."
)

# --- Repository Structure ---
doc.add_heading("Repository Structure", level=1)
doc.add_paragraph(
    "final-project/\n"
    "├─ kusina_app.py                 # CLI entry (previously kusina_bot_clean.py)\n"
    "├─ bot/\n"
    "│  ├─ __init__.py\n"
    "│  ├─ io.py                      # Handles CLI loop, reply language state, voice I/O\n"
    "│  ├─ data.py                    # RecipeDoc model, loaders, Chroma vectorstore binding\n"
    "│  ├─ nlp.py                     # Language detection, translation, ingredient localization\n"
    "│  ├─ tools.py                   # LangChain tools: vector_search, keyword_search, transcribe, summarize, nutrition, shopping list, etc.\n"
    "│  └─ agent.py                   # Builds LangChain agent, defines system prompt, handles routing via chat_once()\n"
    "├─ state/\n"
    "│  └─ kusina_memory.json         # Persistent conversation memory\n"
    "├─ exports/                      # Folder for cookbooks, logs, exported results\n"
    "└─ .env                          # Environment variables (API keys, configs)"
)

# --- Explanation of Tools ---
doc.add_heading("Tools in tools.py", level=1)
doc.add_paragraph(
    "- vector_search: semantic recipe retrieval from Chroma vector DB.\n"
    "- keyword_search: fuzzy fallback search using RapidFuzz.\n"
    "- transcribe_media: transcribes YouTube or local audio using API/Whisper.\n"
    "- estimate_nutrition: estimates calories/macros via LLM prompt.\n"
    "- make_shopping_list: aggregates ingredients into a grouped shopping list.\n"
    "- create_cookbook: exports multiple recipes into a Markdown cookbook.\n"
    "- add_feedback: appends feedback to feedback.jsonl.\n"
    "- translate_text: translates any text into target language.\n"
    "- summarize_video: extracts transcript, summarizes into title, key ingredients, steps, tip.\n"
    "- qa_video: answers a specific question from video transcript.\n"
    "- ingest_link: parses YouTube or article link into structured recipe card.\n"
    "- calories_from_url: extracts ingredients from link, estimates nutrition."
)

# --- Agent Flow ---
doc.add_heading("Agent Flow (agent.py)", level=1)
doc.add_paragraph(
    "1. User sends text or media.\n"
    "2. If translation intent: runs translate_text.\n"
    "3. If input has a URL: routes to ingest_link or calories_from_url.\n"
    "4. If calories requested: checks cached session hits or extracts ingredients.\n"
    "5. Otherwise: forwards to LangChain agent with memory and tools.\n\n"
    "build_agent() sets up:\n"
    "- Tools (from tools.py)\n"
    "- Memory (ConversationBufferMemory)\n"
    "- System prompt (defines tone, workflow, rules)\n"
    "- LangSmith callbacks for monitoring."
)

# --- kusina_app.py and streamlit_app.py ---
doc.add_heading("Application Entry Points", level=1)
doc.add_paragraph(
    "kusina_app.py:\n"
    "- CLI interface for chatbot testing.\n"
    "- Loads vectorstore, session, and agent.\n"
    "- Accepts user input, routes via chat_once().\n"
    "- Prints results to terminal.\n\n"
    "streamlit_app.py:\n"
    "- Web UI using Streamlit.\n"
    "- Lets users enter text, URLs, or audio.\n"
    "- Displays chatbot replies in styled format.\n"
    "- Maintains session memory across interactions.\n"
    "- Enables deployment as web/mobile accessible chatbot."
)

# --- Evaluation ---
doc.add_heading("Evaluation", level=1)
doc.add_paragraph(
    "Success criteria include:\n"
    "- Accuracy of QA responses.\n"
    "- Transcript and retrieval reliability.\n"
    "- Low latency conversational interface.\n"
    "- Proper session memory handling.\n"
    "- Clear documentation and working deployment."
)

# Save the document
output_path = "/mnt/data/Multimodal_AI_ChatBot_Report.docx"
doc.save(output_path)
output_path
